import random
from docx import Document  # 导入相关库
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Cm, Pt, RGBColor
from docx.document import Document as Doc  # 导入这个库的原因是在我们写相关函数的时候会得到提示，要不然许多函数得不到提示无法写出来，太过复杂了

max = 50
num = 3
qNum = 40*20
op_min = 0
op_max = 1

max_ex_len = 0


def cExp(ex=None, left=None):
    if not left:
        left = random.randint(0, max)
    if not ex:
        ex = str(left).ljust(4, ' ')
    if random.randint(op_min, op_max) == 1:
        mid = random.randint(0, left)
        s = ex + '-'.center(5, ' ') + str(mid).ljust(4, ' ')
        v = left - mid
    else:
        mid = random.randint(0, max - left)
        s = ex + '+'.center(5, ' ') + str(mid).ljust(4, ' ')
        v = left + mid
    return [s, v]


qs = []
for i in range(qNum):
    left = random.randint(0, max)
    ex = None
    v = None
    for j in range(num - 1):
        (ex, v) = cExp(ex, v)
    # print(ex, '=', str(v).ljust(2, ' '))
    # print(ex, '=', ''.ljust(2, ' '))
    if max_ex_len < len(ex):
        max_ex_len = len(ex)
    qs.append(ex + ('='.center(4, ' ')))

document = Document()
records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

numbers = ['①', '②', '③', '④', '⑤', '⑥', '⑦', '⑧', '⑨', '⑩', '⑪', '⑫', '⑬', '⑭', '⑮', '⑯', '⑰', '⑱', '⑲', '⑳', '㉑', '㉒',
           '㉓', '㉔', '㉕', '㉖', '㉗', '㉘', '㉙', '㉚', '㉛', '㉜', '㉝', '㉞', '㉟', '㊱', '㊲', '㊳', '㊴', '㊵', '㊶', '㊷', '㊸', '㊹',
           '㊺', '㊻', '㊼', '㊽', '㊾', '㊿']
table = document.add_table(rows=0, cols=4, style='Normal Table')
font_size = 15
num_font_color = RGBColor(242, 138, 121)
font_color = RGBColor(0, 0, 0)
row_height = Cm(1)

for i in range(0, len(qs), 2):
    num = (i % 40)
    row = table.add_row()
    row.height = row_height
    row_cells = row.cells

    row_cells[0].width = Cm(1)
    p = row_cells[0].paragraphs[0]
    run = p.add_run(numbers[num].ljust(3, ' '))
    run.font.color.rgb = num_font_color
    run.font.size = Pt(font_size)
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    row_cells[1].width = Cm(7)
    p = row_cells[1].paragraphs[0]
    run = p.add_run(qs[i])
    run.font.color.rgb = font_color
    run.font.size = Pt(font_size)
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


    num = ((i+1) % 40)
    row_cells[2].width = Cm(1)
    p = row_cells[2].paragraphs[0]
    run = p.add_run(numbers[num].ljust(3, ' '))
    run.font.color.rgb = num_font_color
    run.font.size = Pt(font_size)
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    row_cells[3].width = Cm(7)
    p = row_cells[3].paragraphs[0]
    run = p.add_run(qs[i+1])
    run.font.color.rgb = font_color
    run.font.size = Pt(font_size)
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

# 保存.docx文档
document.save('50-3混合.docx')
